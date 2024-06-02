from docx import Document
import os
from docx.shared import Cm



def main():
    document = Document()

    folder_path = "../output/graphs_new_09_15"

    image_type = "People/activaiton_bar_plots"


    number_of_subjects = 13

    for subject_1_index in range(number_of_subjects):
        for subject_2_index in range(number_of_subjects):
            if subject_1_index >= subject_2_index:
                continue



            file_name = str(subject_1_index) + "_" + str(subject_2_index)

            document.add_heading("subjects " + str(subject_1_index) + " and subject " + str(subject_2_index), 4)

            file_name += ".jpeg"

            image_height = Cm(10)
            image_width = Cm(6.67)
            #Add text:
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            run.add_text("distances_alignment \t\t\t\t\t")

            run_2 = paragraph.add_run()
            run_2.add_text("spearman_alignment")

            #Add images
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(get_image_path(folder_path, "distances_alignment", image_type, file_name), width=image_width, height=image_height)
            run_2 = paragraph.add_run()
            run_2.add_picture(get_image_path(folder_path, "spearman_alignment", image_type, file_name), width=image_width, height=image_height)


    document.save('test.docx')


def get_image_path(folder_path, alignment_name, image_type, file_name):
    return os.path.join(folder_path,
                        alignment_name,
                        "4",
                        "ContactKind.High",
                        image_type,
                        file_name)


if __name__ == "__main__":
    main()

    # folder_path = "../output/graphs_new_09_15"
    # image_type = "People/activaiton_bar_plots"


    # document = Document()
    # document.add_heading("first comparison", 4)
    #
    #
    #
    # # Add text:
    # paragraph = document.add_paragraph()
    # run = paragraph.add_run()
    # run.add_text("dog 1:")
    #
    # run_2 = paragraph.add_run()
    # run_2.add_text("dog 2:")
    #
    # # Add images
    # image_height = Cm(8)
    # image_width = Cm(5)
    #
    # paragraph = document.add_paragraph()
    # run = paragraph.add_run()
    # run.add_picture("./dog1.jpeg", width=image_width,
    #                 height=image_height)
    # run_2 = paragraph.add_run()
    # run_2.add_picture("dog2.jpeg", width=image_width, height=image_height)
    #
    # document.save('test.docx')

